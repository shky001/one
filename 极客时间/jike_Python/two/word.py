# Word 格式
# pip install python-docx

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('标题',0)

# 加粗
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# 标题
document.add_heading('Heading, level 1', level=1)

# 内容

# 表格
records = (
    (3, '101', 'Spam'),
    (7, '422', 'eggs'),
    (4, '631', 'hello, sunshine')
)

table = document.add_table(rows=1, cols=3)

# 表格首行
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

# 保存文件名
document.save('sample.docx')